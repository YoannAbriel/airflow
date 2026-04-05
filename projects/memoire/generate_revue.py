#!/usr/bin/env python3
"""Génère la revue de littérature en DOCX pour le mémoire de Yoann ABRIEL."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# ── Styles ──
style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(12)
style_normal.paragraph_format.line_spacing = 1.5
style_normal.paragraph_format.space_after = Pt(6)

for level in ['Heading 1', 'Heading 2', 'Heading 3']:
    s = doc.styles[level]
    s.font.name = 'Times New Roman'
    s.font.color.rgb = RGBColor(0, 0, 0)

doc.styles['Heading 1'].font.size = Pt(16)
doc.styles['Heading 2'].font.size = Pt(14)
doc.styles['Heading 3'].font.size = Pt(12)

# ── Helper ──
def add_para(text, bold=False, italic=False, style='Normal'):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.first_line_indent = Inches(0.3)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_title_page():
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("REVUE DE LITTÉRATURE")
    r.bold = True
    r.font.size = Pt(20)
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("\nComment l'intelligence artificielle peut-elle transformer\nles processus de veille concurrentielle\ndans les entreprises de télécommunications ?")
    r2.font.size = Pt(14)
    r2.italic = True
    
    for _ in range(4):
        doc.add_paragraph()
    
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run("Yoann ABRIEL\n").bold = True
    p3.add_run("Master 2 — Epitech Digital\n")
    p3.add_run("Spécialisation Data Science & Business Intelligence\n")
    p3.add_run("2025-2026")
    
    doc.add_page_break()

add_title_page()

# ── TABLE DES MATIÈRES (placeholder) ──
doc.add_heading("Table des matières", level=1)
toc_items = [
    "Introduction",
    "1. La veille concurrentielle : fondements théoriques et évolutions",
    "   1.1. Définitions et cadre conceptuel",
    "   1.2. Le cycle du renseignement concurrentiel",
    "   1.3. De la veille manuelle à la veille numérique",
    "   1.4. Enjeux spécifiques dans le secteur des télécommunications",
    "2. L'intelligence artificielle appliquée à la Business Intelligence",
    "   2.1. Le traitement automatique du langage naturel et le text mining",
    "   2.2. L'apprentissage automatique pour la détection de tendances",
    "   2.3. Le web scraping intelligent et la collecte automatisée",
    "3. Les architectures RAG (Retrieval-Augmented Generation)",
    "   3.1. Principes fondamentaux et fonctionnement",
    "   3.2. Applications en contexte entreprise",
    "   3.3. RAG versus fine-tuning : avantages comparatifs pour la veille",
    "4. Les grands modèles de langage et chatbots pour la gestion des connaissances",
    "   4.1. État de l'art des LLMs",
    "   4.2. Chatbots conversationnels en entreprise",
    "   4.3. Génération automatique de synthèses et rapports",
    "5. Cadre éthique, limites et perspectives",
    "   5.1. Biais algorithmiques et équité",
    "   5.2. RGPD, AI Act et protection des données personnelles",
    "   5.3. Hallucinations et fiabilité des LLMs",
    "   5.4. La complémentarité humain-IA",
    "Conclusion",
    "Références bibliographiques",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)

doc.add_page_break()

# ══════════════════════════════════════════════════
# INTRODUCTION
# ══════════════════════════════════════════════════
doc.add_heading("Introduction", level=1)

add_para(
    "Dans un environnement économique marqué par une concurrence accrue et une accélération sans précédent "
    "des cycles d'innovation, la capacité des entreprises à surveiller, analyser et anticiper les mouvements "
    "de leurs concurrents est devenue un facteur déterminant de leur pérennité. Le secteur des télécommunications, "
    "caractérisé par une convergence technologique rapide, des investissements massifs en infrastructures et une "
    "régulation en constante évolution, illustre de manière exemplaire cette réalité. Les opérateurs tels qu'Orange, "
    "Deutsche Telekom ou Vodafone doivent non seulement surveiller les stratégies commerciales de leurs concurrents "
    "directs, mais aussi anticiper les disruptions provenant d'acteurs issus du numérique — les « over-the-top » (OTT) "
    "comme Google, Amazon ou Netflix — qui redéfinissent les contours de l'industrie."
)

add_para(
    "Parallèlement, l'émergence de l'intelligence artificielle (IA), et plus particulièrement des grands modèles "
    "de langage (Large Language Models, LLMs), ouvre des perspectives inédites pour transformer les processus "
    "de veille concurrentielle. L'IA permet désormais d'automatiser la collecte de données massives, d'extraire "
    "des informations pertinentes de sources hétérogènes et non structurées, et de produire des synthèses "
    "analytiques à une vitesse et une échelle inaccessibles aux méthodes traditionnelles. Selon le rapport "
    "McKinsey « The State of AI in 2025 », trente-neuf pour cent des entreprises interrogées attribuent déjà "
    "un impact mesurable de l'IA sur leur résultat opérationnel (McKinsey, 2025). Gartner prévoit quant à lui "
    "que les dépenses mondiales en services d'IA atteindront 478 milliards de dollars d'ici 2028 (Gartner, 2024)."
)

add_para(
    "La présente revue de littérature s'inscrit dans le cadre d'un mémoire de Master 2 dont la problématique "
    "est la suivante : « Comment l'intelligence artificielle peut-elle transformer les processus de veille "
    "concurrentielle dans les entreprises de télécommunications ? ». Le cas d'application porte sur le "
    "développement d'un MVP de veille automatisée pour Orange Business, combinant un chatbot fondé sur "
    "l'architecture RAG (Retrieval-Augmented Generation) et la génération de synthèses hebdomadaires. "
    "Cette revue vise à établir un état de l'art rigoureux en mobilisant cinq axes thématiques : les fondements "
    "théoriques de la veille concurrentielle, l'apport de l'IA à la Business Intelligence, les architectures RAG, "
    "les LLMs appliqués à la gestion des connaissances, et enfin le cadre éthique et les limites de ces technologies."
)

doc.add_page_break()

# ══════════════════════════════════════════════════
# CHAPITRE 1 : VEILLE CONCURRENTIELLE
# ══════════════════════════════════════════════════
doc.add_heading("1. La veille concurrentielle : fondements théoriques et évolutions", level=1)

doc.add_heading("1.1. Définitions et cadre conceptuel", level=2)

add_para(
    "La veille concurrentielle, ou competitive intelligence (CI) dans la littérature anglo-saxonne, désigne "
    "l'ensemble des activités systématiques de collecte, d'analyse et de diffusion d'informations relatives "
    "à l'environnement concurrentiel d'une organisation, dans le but d'éclairer la prise de décision stratégique. "
    "La Society of Competitive Intelligence Professionals (SCIP), fondée en 1986 et devenue l'une des principales "
    "associations professionnelles du domaine, définit la CI comme « la collecte et l'analyse légales d'informations "
    "concernant les capacités, les vulnérabilités et les intentions des concurrents, menées à partir de bases de "
    "données et de sources ouvertes selon des pratiques éthiques » (SCIP, 2000). Cette définition met en lumière "
    "deux dimensions essentielles : le caractère légal et éthique de la démarche, qui distingue la veille "
    "concurrentielle de l'espionnage industriel, et sa finalité décisionnelle."
)

add_para(
    "Michael Porter, dans son ouvrage fondateur « Competitive Advantage » (Porter, 1985), a posé les bases "
    "conceptuelles de l'analyse concurrentielle en proposant le modèle des cinq forces qui structurent la "
    "dynamique d'un secteur : la rivalité entre concurrents existants, la menace de nouveaux entrants, le "
    "pouvoir de négociation des fournisseurs et des clients, et la menace des produits de substitution. Ce "
    "modèle demeure un cadre analytique de référence pour la veille concurrentielle, car il permet de "
    "structurer la collecte d'informations autour de dimensions stratégiques bien identifiées. Porter a "
    "ultérieurement affiné son approche en insistant sur l'importance de comprendre non seulement la "
    "position actuelle des concurrents, mais aussi leurs intentions et leurs capacités futures (Porter, 2008)."
)

add_para(
    "Le concept d'intelligence économique, plus large que celui de veille concurrentielle, a été formalisé "
    "en France notamment par le rapport Martre (1994), qui le définit comme « l'ensemble des actions "
    "coordonnées de recherche, de traitement et de distribution, en vue de son exploitation, de l'information "
    "utile aux acteurs économiques ». Ce rapport a posé les fondements d'une politique publique d'intelligence "
    "économique en France, distinguant trois volets complémentaires : la veille (collecte et analyse), la "
    "protection du patrimoine informationnel, et l'influence (Martre, 1994). Cette tripartition reste pertinente "
    "pour comprendre le positionnement de la veille concurrentielle au sein d'un dispositif plus global "
    "d'intelligence stratégique."
)

doc.add_heading("1.2. Le cycle du renseignement concurrentiel", level=2)

add_para(
    "La majorité des chercheurs conceptualisent le processus de veille concurrentielle comme un cycle "
    "itératif composé de phases interdépendantes (Pellissier et Nenzhelele, 2013). Ce modèle, largement "
    "inspiré du cycle du renseignement utilisé par les agences de sécurité nationale, notamment la CIA, "
    "comprend généralement quatre à six étapes : la planification et l'orientation des besoins, la collecte "
    "des données, le traitement et l'analyse, et enfin la diffusion des résultats aux décideurs. "
    "Pellissier et Nenzhelele (2013), dans leur tentative de proposer un modèle universel du processus de CI, "
    "ont montré que malgré des variations terminologiques, la structure cyclique du processus fait consensus "
    "dans la littérature académique."
)

add_para(
    "Le modèle PCMAC (Plan, Capture, Manage, Analyse, Communicate), décrit dans la littérature spécialisée "
    "et repris par plusieurs auteurs sur la plateforme ScienceDirect, insiste sur l'importance cruciale de la "
    "première phase — la planification et la priorisation — pour le succès global de l'activité de veille. "
    "En effet, sans une définition précise des besoins informationnels et des questions stratégiques à résoudre, "
    "le risque est de collecter un volume considérable de données sans parvenir à en extraire des insights "
    "actionnables. Cette phase d'orientation, souvent négligée dans les organisations, conditionne la pertinence "
    "de l'ensemble du cycle (Trim et Lee, 2008). Rouach et Santi (2001) soulignent également que la veille "
    "concurrentielle ne se limite pas à la surveillance réactive des concurrents, mais doit s'inscrire dans "
    "une démarche proactive d'anticipation des ruptures stratégiques."
)

doc.add_heading("1.3. De la veille manuelle à la veille numérique", level=2)

add_para(
    "L'histoire de la veille concurrentielle révèle une transformation profonde des méthodes et des outils "
    "mobilisés. Dans les années 1980 et 1990, la veille reposait principalement sur des sources documentaires "
    "traditionnelles — presse spécialisée, rapports annuels, brevets, études de marché — et sur des réseaux "
    "humains de collecte informelle (salons professionnels, contacts sectoriels). Le rapport du CIGREF de 1998 "
    "sur la veille stratégique témoigne de cette époque où l'enjeu principal résidait dans l'accès à l'information "
    "plutôt que dans sa surabondance (CIGREF, 1998)."
)

add_para(
    "L'avènement d'Internet et la prolifération des sources numériques ont radicalement modifié la donne. "
    "La veille concurrentielle dispose désormais d'un volume colossal de données accessibles en temps réel : "
    "sites web des concurrents, réseaux sociaux, flux RSS, forums spécialisés, bases de données de brevets "
    "en ligne, communiqués de presse. Cette abondance informationnelle, si elle constitue une opportunité "
    "considérable, pose également le défi de la surcharge informationnelle (information overload) et de la "
    "nécessité de disposer d'outils capables de filtrer, trier et hiérarchiser les données pertinentes. "
    "C'est précisément dans ce contexte que les technologies d'intelligence artificielle trouvent leur "
    "pertinence, en permettant d'automatiser des tâches jusqu'alors manuelles et chronophages."
)

doc.add_heading("1.4. Enjeux spécifiques dans le secteur des télécommunications", level=2)

add_para(
    "Le secteur des télécommunications présente des caractéristiques qui rendent la veille concurrentielle "
    "à la fois critique et particulièrement complexe. En premier lieu, l'intensité concurrentielle y est "
    "élevée, avec un nombre limité d'acteurs majeurs par marché national qui se disputent des parts de marché "
    "dans un contexte de saturation des abonnements mobiles dans les pays développés. La convergence entre "
    "télécommunications, médias et technologies de l'information brouille les frontières sectorielles "
    "traditionnelles et multiplie le nombre d'acteurs à surveiller."
)

add_para(
    "En second lieu, le rythme d'innovation technologique — déploiement de la 5G, développement de l'edge "
    "computing, émergence des réseaux privés pour les entreprises — impose une veille technologique soutenue. "
    "Les opérateurs doivent anticiper les évolutions des standards techniques, les stratégies d'investissement "
    "de leurs concurrents dans les infrastructures, et les nouveaux usages rendus possibles par ces technologies. "
    "Pour un acteur comme Orange Business, qui opère sur le segment B2B des services numériques aux entreprises, "
    "la veille doit couvrir un spectre particulièrement large : concurrents directs (opérateurs télécoms), "
    "acteurs du cloud (AWS, Azure, Google Cloud), intégrateurs de solutions IT, et pure players de la cybersécurité. "
    "La capacité à synthétiser ces informations dispersées en insights actionnables constitue un avantage "
    "concurrentiel déterminant."
)

add_para(
    "Enfin, la dimension réglementaire est omniprésente dans les télécommunications. Les décisions des "
    "autorités de régulation nationales (ARCEP en France, Bundesnetzagentur en Allemagne) et européennes "
    "concernant l'attribution des fréquences, la régulation des tarifs d'interconnexion ou les obligations "
    "de couverture territoriale peuvent modifier substantiellement l'équilibre concurrentiel. La veille "
    "réglementaire constitue donc un volet essentiel du dispositif de veille des opérateurs."
)

doc.add_page_break()

# ══════════════════════════════════════════════════
# CHAPITRE 2 : IA APPLIQUÉE À LA BI
# ══════════════════════════════════════════════════
doc.add_heading("2. L'intelligence artificielle appliquée à la Business Intelligence", level=1)

doc.add_heading("2.1. Le traitement automatique du langage naturel et le text mining", level=2)

add_para(
    "Le traitement automatique du langage naturel (Natural Language Processing, NLP) constitue l'une des "
    "branches de l'intelligence artificielle les plus directement pertinentes pour la veille concurrentielle. "
    "Le NLP vise à doter les machines de la capacité de comprendre, interpréter et générer du langage humain. "
    "Appliqué à la veille, il permet l'extraction automatisée d'informations structurées à partir de sources "
    "textuelles non structurées — articles de presse, rapports financiers, publications sur les réseaux sociaux, "
    "brevets — qui constituent la matière première de l'analyste de veille."
)

add_para(
    "Le text mining, défini comme l'ensemble des techniques permettant d'extraire des connaissances utiles "
    "à partir de grandes collections de textes, s'appuie sur un éventail de méthodes issues du NLP : la "
    "reconnaissance d'entités nommées (Named Entity Recognition, NER) pour identifier les acteurs, les "
    "produits et les lieux mentionnés dans les documents ; l'analyse de sentiment (sentiment analysis) pour "
    "évaluer la tonalité des discours publics sur les concurrents ; la classification automatique de documents "
    "par thématiques ; et le résumé automatique (text summarization) pour condenser de grands volumes de textes. "
    "Une revue systématique de la littérature publiée par Sahut et al. (2021) dans la revue International "
    "Journal of Information Management a montré que l'analyse des médias sociaux, l'analyse de marché et "
    "la veille concurrentielle constituent les thématiques dominantes des applications du text mining dans "
    "le domaine du management des services."
)

add_para(
    "L'avènement des architectures Transformer, introduites par Vaswani et al. (2017) dans l'article "
    "fondateur « Attention Is All You Need », a constitué une rupture paradigmatique dans le domaine du NLP. "
    "En remplaçant les mécanismes de récurrence par des mécanismes d'attention qui permettent de traiter "
    "les séquences textuelles de manière parallèle, les Transformers ont considérablement amélioré les "
    "performances des modèles de langage sur l'ensemble des tâches du NLP. Le modèle BERT (Bidirectional "
    "Encoder Representations from Transformers), proposé par Devlin et al. (2019), a démontré la puissance "
    "du pré-entraînement bidirectionnel pour les tâches de compréhension du langage, tandis que la série "
    "des modèles GPT a exploré la voie de la génération de texte (Brown et al., 2020). Ces avancées ont "
    "rendu possibles des applications de veille concurrentielle impensables quelques années auparavant, "
    "comme l'extraction automatique de signaux faibles à partir de flux d'informations continus."
)

doc.add_heading("2.2. L'apprentissage automatique pour la détection de tendances", level=2)

add_para(
    "Au-delà du NLP, les techniques d'apprentissage automatique (machine learning) offrent des capacités "
    "analytiques puissantes pour la détection de patterns et de tendances dans les données de veille. "
    "Les algorithmes de clustering permettent de regrouper automatiquement des documents ou des signaux "
    "thématiquement proches, facilitant ainsi l'identification de tendances émergentes. Les méthodes de "
    "détection d'anomalies peuvent signaler des événements inhabituels — un pic soudain de recrutements "
    "chez un concurrent, une variation significative de sa politique tarifaire — qui méritent une attention "
    "particulière de la part des analystes."
)

add_para(
    "L'analyse de séries temporelles, enrichie par des techniques de deep learning telles que les réseaux "
    "LSTM (Long Short-Term Memory) et les Transformers temporels, permet de modéliser l'évolution des "
    "indicateurs concurrentiels et d'identifier des ruptures de tendance. Dans le secteur des télécommunications, "
    "ces approches trouvent des applications concrètes dans le suivi en temps réel des parts de marché, "
    "l'analyse des stratégies tarifaires des concurrents, ou encore la prédiction des lancements de produits "
    "à partir de signaux précurseurs tels que les dépôts de brevets ou les offres d'emploi. Le rapport "
    "de l'OCDE sur l'intelligence artificielle et les dynamiques concurrentielles (OCDE, 2025) souligne "
    "que l'adoption de l'IA par les entreprises modifie en profondeur les dynamiques de marché, les firmes "
    "les plus avancées en matière d'IA disposant d'un avantage informationnel significatif."
)

doc.add_heading("2.3. Le web scraping intelligent et la collecte automatisée", level=2)

add_para(
    "La collecte automatisée de données sur le web, ou web scraping, constitue la première étape de toute "
    "chaîne de veille numérique. Les techniques de web scraping ont considérablement évolué ces dernières "
    "années, passant de scripts rudimentaires d'extraction HTML à des systèmes intelligents capables de "
    "naviguer dans des sites web dynamiques, de contourner les mécanismes anti-robots et d'adapter leur "
    "comportement aux changements de structure des pages. L'intégration de l'apprentissage automatique dans "
    "les outils de web scraping a permis l'émergence de ce que l'on peut qualifier de « scraping intelligent » : "
    "des agents capables de comprendre le contexte d'une page, d'identifier automatiquement les éléments "
    "pertinents à extraire (noms de produits, prix, caractéristiques techniques) et de s'adapter aux "
    "évolutions des sites cibles sans intervention humaine."
)

add_para(
    "Dans le contexte de la veille concurrentielle, le web scraping automatisé permet une surveillance "
    "continue des sites web des concurrents, des places de marché, des portails d'offres d'emploi et "
    "des bases de données réglementaires. Des plateformes telles que Browse AI ou Oxylabs illustrent cette "
    "tendance à l'automatisation intelligente de la collecte de données web, en proposant des fonctionnalités "
    "de monitoring en temps réel et de détection de changements. Néanmoins, le web scraping soulève des "
    "questions juridiques et éthiques importantes, notamment en matière de respect des conditions d'utilisation "
    "des sites web, de protection des données personnelles et de loyauté concurrentielle. Ces enjeux "
    "doivent être intégrés dans la conception de tout système de veille automatisée."
)

doc.add_page_break()

# ══════════════════════════════════════════════════
# CHAPITRE 3 : RAG
# ══════════════════════════════════════════════════
doc.add_heading("3. Les architectures RAG (Retrieval-Augmented Generation)", level=1)

doc.add_heading("3.1. Principes fondamentaux et fonctionnement", level=2)

add_para(
    "Le concept de Retrieval-Augmented Generation (RAG), ou génération augmentée par la recherche "
    "d'information, a été formalisé par Lewis et al. (2020) dans un article séminal présenté à la "
    "conférence NeurIPS. Les auteurs y proposent un cadre hybride qui combine un composant de recherche "
    "documentaire (retriever) avec un modèle de génération de texte (generator), permettant ainsi au système "
    "d'ancrer ses réponses dans des documents factuels extraits d'une base de connaissances. Le principe "
    "fondamental du RAG repose sur l'intuition que les modèles de langage purement paramétriques — c'est-à-dire "
    "qui s'appuient uniquement sur les connaissances encodées dans leurs paramètres lors de l'entraînement — "
    "présentent des limites structurelles : ils ne peuvent pas facilement mettre à jour leurs connaissances, "
    "ils peinent à fournir des explications sur l'origine de leurs réponses, et ils sont sujets aux "
    "hallucinations (Lewis et al., 2020)."
)

add_para(
    "Le fonctionnement technique d'une architecture RAG se décompose en plusieurs étapes. Premièrement, "
    "les documents de la base de connaissances sont découpés en segments (chunks) puis transformés en "
    "représentations vectorielles (embeddings) à l'aide d'un modèle d'encodage sémantique. Ces embeddings "
    "sont stockés dans une base de données vectorielle (vector store) telle que Pinecone, Weaviate ou "
    "ChromaDB. Lorsqu'un utilisateur soumet une requête, celle-ci est à son tour transformée en embedding, "
    "et une recherche par similarité est effectuée dans la base vectorielle pour retrouver les documents "
    "les plus pertinents. Ces documents sont alors injectés dans le prompt du modèle de langage, qui "
    "génère une réponse contextualisée en s'appuyant sur les informations retrouvées. Ce processus est "
    "parfois qualifié de mémoire non paramétrique, par opposition à la mémoire paramétrique intrinsèque "
    "au modèle (Springer, 2025)."
)

add_para(
    "Depuis l'article fondateur de Lewis et al., l'architecture RAG a fait l'objet de nombreuses "
    "améliorations. Les surveys récents (Gao et al., 2024 ; Singh et Ehtesham, 2025) identifient "
    "plusieurs axes d'évolution : l'amélioration des stratégies de découpage des documents, le "
    "raffinement des méthodes de recherche sémantique (hybrid search combinant recherche vectorielle "
    "et lexicale), le re-ranking des documents retrouvés pour en améliorer la pertinence, et l'émergence "
    "du concept de « RAG agentique » où des agents autonomes orchestrent dynamiquement le processus "
    "de recherche et de génération. Le cadre R2AG proposé par Ye et al. (2024) illustre cette tendance "
    "en introduisant un re-ranking récursif des documents pendant la génération, adaptant dynamiquement "
    "les priorités en fonction de l'évolution de la réponse."
)

doc.add_heading("3.2. Applications en contexte entreprise", level=2)

add_para(
    "L'adoption du RAG en entreprise s'est considérablement accélérée depuis 2023, portée par la "
    "démocratisation des LLMs et la nécessité de les ancrer dans les connaissances propres à chaque "
    "organisation. L'un des cas d'usage les plus documentés est celui de Morgan Stanley, qui a déployé "
    "un chatbot interne alimenté par GPT-4 et connecté à sa vaste base documentaire de recherche "
    "financière, permettant à ses conseillers en gestion de patrimoine d'accéder instantanément à des "
    "synthèses pertinentes issues de milliers de rapports internes (ZenML, 2024). Ce déploiement illustre "
    "le potentiel du RAG pour la gestion des connaissances en entreprise : plutôt que de réentraîner un "
    "modèle coûteusement sur des données propriétaires, le RAG permet d'injecter dynamiquement le contexte "
    "organisationnel dans les réponses du modèle."
)

add_para(
    "Miyaji et al. (2025), dans un article présenté à la conférence IEEE on Artificial Intelligence, "
    "démontrent comment un système RAG avancé peut transformer la prise de décision en entreprise en "
    "permettant une interrogation en langage naturel des bases de connaissances organisationnelles. "
    "L'étude de Veturi et al. (2024), citée dans une revue publiée par Springer dans la revue Business "
    "& Information Systems Engineering, montre comment l'utilisation de données organisationnelles "
    "(documents de politique interne, procédures, guides) dans un système RAG améliore significativement "
    "la performance du système par rapport à un LLM générique. Ces résultats convergent vers un constat : "
    "le RAG constitue aujourd'hui la méthode privilégiée pour adapter les LLMs à des contextes "
    "organisationnels spécifiques sans les coûts et la complexité du fine-tuning."
)

doc.add_heading("3.3. RAG versus fine-tuning : avantages comparatifs pour la veille", level=2)

add_para(
    "Le choix entre RAG et fine-tuning constitue une décision architecturale structurante pour tout "
    "projet d'IA appliquée à la veille concurrentielle. Le fine-tuning consiste à poursuivre l'entraînement "
    "d'un modèle pré-entraîné sur un jeu de données spécifique au domaine cible, modifiant ainsi les "
    "paramètres internes du modèle pour qu'il intègre des connaissances spécialisées. Le RAG, en revanche, "
    "laisse le modèle inchangé et lui fournit le contexte pertinent au moment de l'inférence. Selon IBM "
    "(2025), la différence fondamentale réside dans le fait que le RAG connecte le LLM à une base de "
    "données propriétaire, tandis que le fine-tuning optimise le modèle pour des tâches spécifiques "
    "à un domaine."
)

add_para(
    "Pour un système de veille concurrentielle, le RAG présente plusieurs avantages décisifs. Premièrement, "
    "la fraîcheur de l'information : les données de veille évoluent quotidiennement, et le RAG permet "
    "d'intégrer immédiatement de nouveaux documents sans réentraînement du modèle. Deuxièmement, la "
    "traçabilité : le RAG permet de citer les sources documentaires sur lesquelles repose chaque réponse, "
    "un impératif en matière de veille où la crédibilité de l'information dépend de sa provenance. "
    "Troisièmement, le coût : le fine-tuning d'un LLM de grande taille requiert des ressources "
    "computationnelles considérables et une expertise technique pointue, tandis que le RAG peut être "
    "implémenté avec des ressources plus modestes. Comme le souligne Oracle (2024), le RAG est plus "
    "simple à implémenter tout en offrant une flexibilité supérieure pour les environnements dynamiques."
)

add_para(
    "Toutefois, le fine-tuning conserve sa pertinence dans certains cas. Pour des domaines stables "
    "dont les connaissances évoluent peu, le fine-tuning peut offrir de meilleures performances sans "
    "la surcharge liée à la recherche documentaire en temps réel. En pratique, plusieurs études suggèrent "
    "que la combinaison des deux approches — un modèle fine-tuné augmenté par du RAG — peut produire "
    "des résultats optimaux, comme le montrent les travaux comparatifs sur les modèles FlanT5 (Khadka, 2024). "
    "Pour un système de veille concurrentielle dans les télécommunications, où l'environnement informationnel "
    "est hautement dynamique, le RAG apparaît néanmoins comme l'approche la plus adaptée."
)

doc.add_page_break()

# ══════════════════════════════════════════════════
# CHAPITRE 4 : LLMs ET CHATBOTS
# ══════════════════════════════════════════════════
doc.add_heading("4. Les grands modèles de langage et chatbots pour la gestion des connaissances", level=1)

doc.add_heading("4.1. État de l'art des LLMs", level=2)

add_para(
    "Les grands modèles de langage (Large Language Models, LLMs) ont connu une progression fulgurante "
    "depuis l'introduction de l'architecture Transformer par Vaswani et al. (2017). Le modèle GPT-3, "
    "présenté par Brown et al. (2020) avec ses 175 milliards de paramètres, a démontré la capacité des "
    "LLMs à réaliser un éventail remarquable de tâches linguistiques en « few-shot learning », c'est-à-dire "
    "à partir d'un nombre très réduit d'exemples fournis dans le prompt. Cette démonstration a ouvert "
    "la voie à une course à l'échelle qui a vu émerger des modèles toujours plus performants : GPT-4 "
    "(OpenAI, 2023), Claude 3 puis Claude 3.5 (Anthropic, 2024), Gemini (Google, 2024), et LLaMA "
    "(Meta, 2023) dans le domaine des modèles ouverts."
)

add_para(
    "L'état de l'art des LLMs en 2025 se caractérise par plusieurs tendances majeures. D'une part, "
    "l'augmentation des fenêtres de contexte — passant de quelques milliers de tokens à plusieurs "
    "centaines de milliers, voire des millions — permet aux modèles de traiter des documents longs dans "
    "leur intégralité. D'autre part, l'émergence de capacités de raisonnement améliorées, comme en "
    "témoignent les modèles dits « de raisonnement » (o1 d'OpenAI, Claude avec le mode de pensée "
    "étendu d'Anthropic), ouvre des perspectives pour des tâches analytiques complexes telles que la "
    "synthèse de veille concurrentielle. Enfin, la tendance à la spécialisation sectorielle et à la "
    "réduction de taille des modèles (small language models) rend ces technologies accessibles à un "
    "nombre croissant d'organisations, y compris celles disposant de ressources computationnelles limitées."
)

doc.add_heading("4.2. Chatbots conversationnels en entreprise", level=2)

add_para(
    "Le déploiement de chatbots conversationnels alimentés par des LLMs en entreprise constitue l'une "
    "des applications les plus visibles de l'IA générative. Ces systèmes diffèrent fondamentalement des "
    "chatbots traditionnels à base de règles (rule-based) par leur capacité à comprendre des requêtes "
    "formulées en langage naturel et à générer des réponses contextualisées et fluides. L'article de "
    "Sridhara et al. (2024), publié sur arXiv sous le titre « FACTS About Building Retrieval Augmented "
    "Generation-based Chatbots », souligne que la construction d'un chatbot d'entreprise réussi, même "
    "à l'ère post-ChatGPT, reste un défi d'ingénierie considérable. Le processus exige une ingénierie "
    "méticuleuse des pipelines RAG, un ajustement fin des LLMs, une conception soignée des prompts, "
    "tout en garantissant la pertinence et l'exactitude des connaissances de l'entreprise, le respect "
    "des droits d'accès aux documents et la protection des informations personnelles."
)

add_para(
    "Pour la veille concurrentielle, un chatbot RAG offre un mode d'accès inédit à l'intelligence "
    "concurrentielle. Plutôt que de consulter des tableaux de bord statiques ou de parcourir des rapports "
    "volumineux, les analystes et les décideurs peuvent interroger le système en langage naturel : "
    "« Quelles sont les dernières annonces de Deutsche Telekom concernant la 5G privée ? », « Comment "
    "évolue la stratégie cloud de Vodafone par rapport au trimestre précédent ? ». Le chatbot, grâce "
    "à l'architecture RAG, recherche dans la base documentaire de veille les informations les plus "
    "pertinentes et génère une réponse synthétique assortie de références. Ce mode d'interaction "
    "conversationnel réduit considérablement le temps d'accès à l'information et démocratise l'usage "
    "de la veille concurrentielle au sein de l'organisation."
)

doc.add_heading("4.3. Génération automatique de synthèses et rapports", level=2)

add_para(
    "La génération automatique de synthèses constitue une application particulièrement prometteuse des "
    "LLMs pour la veille concurrentielle. Traditionnellement, la production de rapports de veille "
    "hebdomadaires ou mensuels mobilise un temps considérable de la part des analystes, qui doivent "
    "compiler, recouper et synthétiser des dizaines, voire des centaines de sources. Les LLMs, combinés "
    "à une architecture RAG, permettent d'automatiser une grande partie de ce processus : le système "
    "collecte automatiquement les nouvelles informations sur une période donnée, les classe par "
    "thématique et par importance, et génère une synthèse structurée qui met en lumière les faits "
    "saillants, les tendances émergentes et les signaux faibles."
)

add_para(
    "La revue systématique de la littérature sur le RAG et les LLMs pour la gestion des connaissances "
    "en entreprise, publiée en décembre 2025 sur Preprints.org, analyse 77 études primaires et confirme "
    "le potentiel de ces technologies pour répondre aux défis pratiques des organisations en matière de "
    "gestion documentaire et d'automatisation des rapports (Preprints, 2025). Toutefois, les auteurs "
    "soulignent que l'adoption en entreprise reste freinée par plusieurs facteurs : la qualité variable "
    "des données d'entrée, la nécessité de valider les outputs générés, et les préoccupations relatives "
    "à la confidentialité des données. La génération automatique de synthèses ne vise donc pas à "
    "remplacer l'analyste humain, mais à lui fournir une première version enrichie qu'il pourra "
    "réviser, compléter et valider avant diffusion."
)

doc.add_page_break()

# ══════════════════════════════════════════════════
# CHAPITRE 5 : ÉTHIQUE ET LIMITES
# ══════════════════════════════════════════════════
doc.add_heading("5. Cadre éthique, limites et perspectives", level=1)

doc.add_heading("5.1. Biais algorithmiques et équité", level=2)

add_para(
    "L'utilisation de l'intelligence artificielle dans les processus de veille concurrentielle n'est pas "
    "exempte de risques, au premier rang desquels figurent les biais algorithmiques. Les LLMs, entraînés "
    "sur de vastes corpus de textes issus d'Internet, héritent inévitablement des biais présents dans "
    "leurs données d'entraînement. Comme le souligne la survey exhaustive d'Alansari et Luqman (2025) "
    "publiée dans les archives arXiv, les données utilisées pour l'entraînement des LLMs constituent "
    "une source significative de biais que les modèles peuvent involontairement acquérir et propager. "
    "Ces biais peuvent se manifester de multiples manières : surreprésentation de certaines perspectives "
    "culturelles ou économiques, sous-représentation de certaines régions géographiques ou de certains "
    "acteurs économiques, ou encore reproduction de stéréotypes sectoriels."
)

add_para(
    "Dans le contexte de la veille concurrentielle, les biais algorithmiques peuvent conduire à des "
    "analyses déformées de l'environnement concurrentiel. Un système de veille biaisé pourrait, par "
    "exemple, surestimer systématiquement l'importance de certains concurrents au détriment d'autres, "
    "ou privilégier certaines sources d'information au détriment de sources tout aussi pertinentes mais "
    "moins représentées dans les données d'entraînement du modèle. La prise de conscience de ces biais "
    "et la mise en place de mécanismes de détection et de correction constituent un impératif pour "
    "garantir la fiabilité des systèmes de veille augmentés par l'IA. La survey de Huang et al. (2024), "
    "publiée dans ACM Transactions on Information Systems, propose une taxonomie détaillée des "
    "hallucinations et des biais dans les LLMs, offrant un cadre analytique utile pour les praticiens."
)

doc.add_heading("5.2. RGPD, AI Act et protection des données personnelles", level=2)

add_para(
    "Le cadre juridique européen impose des contraintes strictes en matière de protection des données "
    "personnelles qui doivent être intégrées dès la conception de tout système de veille augmenté par l'IA. "
    "Le Règlement Général sur la Protection des Données (RGPD), entré en vigueur en 2018, s'applique "
    "pleinement aux systèmes d'IA qui traitent des données personnelles. La Commission Nationale de "
    "l'Informatique et des Libertés (CNIL) a publié en janvier 2024 une recommandation spécifique sur "
    "l'application du RGPD au développement des systèmes d'intelligence artificielle (CNIL, 2024), "
    "rappelant les obligations en matière de finalité du traitement, de minimisation des données, de "
    "transparence et de droit à l'explication."
)

add_para(
    "Le Règlement européen sur l'Intelligence Artificielle (AI Act), publié le 12 juillet 2024 et "
    "entrant progressivement en application à partir du 1er août 2024, ajoute une couche réglementaire "
    "supplémentaire. Comme le précise la CNIL dans ses premiers éléments de réponse (CNIL, 2024), "
    "le Règlement sur l'IA ne remplace pas les exigences du RGPD mais vise à les compléter en posant "
    "les conditions requises pour développer et déployer des systèmes d'IA de confiance. Pour un système "
    "de veille concurrentielle, ces contraintes se traduisent concrètement par la nécessité de s'assurer "
    "que les données collectées par web scraping ne contiennent pas de données personnelles traitées "
    "illicitement, que les bases documentaires alimentant le RAG respectent les droits de propriété "
    "intellectuelle, et que les utilisateurs du système sont informés de la nature automatisée des "
    "synthèses produites."
)

doc.add_heading("5.3. Hallucinations et fiabilité des LLMs", level=2)

add_para(
    "Le phénomène des hallucinations constitue l'une des limites les plus préoccupantes des LLMs dans "
    "un contexte de veille concurrentielle où l'exactitude factuelle est primordiale. Une hallucination "
    "se produit lorsque le modèle génère des informations plausibles mais factuellement incorrectes, "
    "inventées de toutes pièces ou en contradiction avec les sources disponibles. La survey de Huang "
    "et al. (2024) distingue trois types d'hallucinations : l'infidélité aux données d'entraînement, "
    "l'infidélité aux données d'entrée (le prompt), et l'infidélité aux faits réels du monde. Dans le "
    "domaine juridique, Dahl et al. (2024) ont documenté la persistance de ces hallucinations même dans "
    "les systèmes RAG, soulignant que la recherche documentaire ne constitue pas une panacée."
)

add_para(
    "Pour un système de veille concurrentielle, les conséquences d'une hallucination peuvent être "
    "significatives : une information fabricée sur la stratégie d'un concurrent pourrait conduire à des "
    "décisions stratégiques erronées. L'architecture RAG réduit substantiellement le risque d'hallucination "
    "en ancrant les réponses dans des documents factuels, mais ne l'élimine pas totalement. Le modèle "
    "peut encore halluciner en interprétant incorrectement les documents retrouvés ou en comblant des "
    "lacunes informationnelles par des inférences non fondées. Des approches telles que le « Confidence-"
    "Calibrated RAG » proposé par Ozaki et al. (2025), qui calibre le niveau de confiance des réponses "
    "en fonction de la qualité des documents retrouvés, ouvrent des pistes prometteuses pour améliorer "
    "la fiabilité des systèmes. En pratique, la mise en place d'un circuit de validation humaine des "
    "outputs critiques reste indispensable."
)

doc.add_heading("5.4. La complémentarité humain-IA", level=2)

add_para(
    "La question de la complémentarité entre l'intelligence artificielle et l'expertise humaine se "
    "pose avec acuité dans le domaine de la veille concurrentielle. Loin de l'opposition simpliste entre "
    "automatisation totale et intervention humaine exclusive, la littérature récente plaide pour une "
    "approche d'intelligence augmentée (augmented intelligence) où l'IA amplifie les capacités de "
    "l'analyste plutôt qu'elle ne le remplace. L'IA excelle dans le traitement de volumes massifs de "
    "données, la détection de patterns récurrents et la synthèse d'informations dispersées. L'humain, "
    "en revanche, conserve un avantage irremplaçable dans l'interprétation contextuelle, le jugement "
    "stratégique, la prise en compte de facteurs intangibles (culture d'entreprise, dynamiques "
    "interpersonnelles, intuitions de marché) et la validation critique des outputs de l'IA."
)

add_para(
    "Dans le cadre d'un système de veille augmenté par l'IA tel que le MVP développé pour Orange Business, "
    "cette complémentarité se traduit par un workflow où l'IA assure la collecte, le prétraitement et la "
    "première synthèse des informations, tandis que l'analyste humain intervient pour valider, enrichir "
    "et contextualiser les insights produits. Les synthèses hebdomadaires générées automatiquement ne "
    "sont pas diffusées telles quelles, mais constituent une matière première que l'expert de veille "
    "transforme en recommandations stratégiques actionnables. Cette approche « human-in-the-loop » "
    "permet de tirer parti de la puissance de calcul de l'IA tout en préservant la rigueur analytique "
    "et le discernement que seul l'humain peut apporter. Elle constitue, selon plusieurs études, "
    "le modèle le plus réaliste et le plus efficace d'intégration de l'IA dans les processus "
    "décisionnels des organisations (McKinsey, 2025)."
)

doc.add_page_break()

# ══════════════════════════════════════════════════
# CONCLUSION
# ══════════════════════════════════════════════════
doc.add_heading("Conclusion", level=1)

add_para(
    "Cette revue de littérature a permis de dresser un panorama des connaissances académiques et "
    "professionnelles à l'intersection de la veille concurrentielle et de l'intelligence artificielle, "
    "dans le contexte spécifique du secteur des télécommunications. Plusieurs enseignements majeurs "
    "se dégagent de cette analyse."
)

add_para(
    "Premièrement, la veille concurrentielle, ancrée dans une tradition théorique riche depuis les "
    "travaux fondateurs de Porter (1985) et formalisée par des modèles processuels tels que le cycle "
    "du renseignement, est confrontée à un défi de transformation numérique qui remet en question ses "
    "méthodes traditionnelles. L'explosion des sources d'information numériques et la vitesse croissante "
    "des dynamiques concurrentielles, particulièrement dans un secteur aussi turbulent que les "
    "télécommunications, rendent indispensable le recours à des outils d'automatisation et d'analyse "
    "avancés."
)

add_para(
    "Deuxièmement, les technologies d'intelligence artificielle — NLP, text mining, machine learning, "
    "web scraping intelligent — offrent des capacités opérationnelles directement applicables à chaque "
    "étape du cycle de veille, de la collecte à la diffusion en passant par l'analyse. Les architectures "
    "RAG, en particulier, se distinguent comme la solution la plus adaptée pour intégrer les LLMs dans "
    "un contexte organisationnel, en permettant d'ancrer les réponses générées dans des bases de "
    "connaissances propriétaires continuellement mises à jour."
)

add_para(
    "Troisièmement, le déploiement de chatbots conversationnels et de systèmes de génération "
    "automatique de synthèses transforme le mode d'accès à l'intelligence concurrentielle, en "
    "le rendant plus rapide, plus intuitif et plus démocratique au sein des organisations. Toutefois, "
    "les défis éthiques et techniques — hallucinations, biais, conformité réglementaire — imposent "
    "la mise en place de garde-fous robustes et le maintien d'une supervision humaine qualifiée."
)

add_para(
    "Ces constats fondent la pertinence de la problématique de ce mémoire et orientent les choix "
    "architecturaux du MVP de veille automatisée développé pour Orange Business. La suite de ce "
    "travail s'attachera à décrire la méthodologie de conception du système, à présenter les résultats "
    "de son déploiement expérimental, et à en évaluer les apports au regard des attentes des utilisateurs "
    "et des objectifs stratégiques de l'organisation."
)

doc.add_page_break()

# ══════════════════════════════════════════════════
# RÉFÉRENCES BIBLIOGRAPHIQUES
# ══════════════════════════════════════════════════
doc.add_heading("Références bibliographiques", level=1)

references = [
    "Alansari, A. et Luqman, H. (2025). Large Language Models Hallucination: A Comprehensive Survey. arXiv preprint arXiv:2510.06265.",
    "Brown, T. B. et al. (2020). Language Models are Few-Shot Learners. Advances in Neural Information Processing Systems, 33, 1877-1901.",
    "CIGREF (1998). Veille stratégique : rapport du groupe de travail. Paris : CIGREF.",
    "CNIL (2024). Délibération n° 2024-011 du 18 janvier 2024 portant adoption d'une recommandation sur l'application du RGPD au développement des systèmes d'intelligence artificielle. Paris : CNIL.",
    "CNIL (2024). Entrée en vigueur du règlement européen sur l'IA : les premières questions-réponses. Paris : CNIL.",
    "Devlin, J. et al. (2019). BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding. Proceedings of NAACL-HLT, 4171-4186.",
    "Gao, Y. et al. (2024). Retrieval-Augmented Generation: A Comprehensive Survey of Architectures, Enhancements, and Robustness Frontiers. arXiv preprint.",
    "Gartner (2024). Market Opportunity Map: Artificial Intelligence Services, Worldwide, 2024. Stamford : Gartner Inc.",
    "Gartner (2024). Hype Cycle for Artificial Intelligence, 2024. Stamford : Gartner Inc.",
    "Huang, L. et al. (2024). A Survey on Hallucination in Large Language Models: Principles, Taxonomy, Challenges, and Open Questions. ACM Transactions on Information Systems.",
    "IBM (2025). RAG vs. Fine-tuning. IBM Think Topics.",
    "Lewis, P. et al. (2020). Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. Advances in Neural Information Processing Systems, 33, 9459-9474.",
    "Martre, H. (1994). Intelligence économique et stratégie des entreprises. Rapport du Commissariat Général du Plan. Paris : La Documentation Française.",
    "McKinsey & Company (2025). The State of AI in 2025: Agents, Innovation, and Transformation. New York : McKinsey Global Institute.",
    "Miyaji, R. et al. (2025). Empowering Business Decisions and Knowledge Management Through Advanced RAG-Driven QA Systems. IEEE Conference on Artificial Intelligence (CAI), 55-60.",
    "OCDE (2025). Artificial Intelligence and Competitive Dynamics in Downstream Markets. Paris : Éditions OCDE.",
    "Oracle (2024). RAG vs. Fine-Tuning: How to Choose. Oracle Cloud Infrastructure.",
    "Ozaki et al. (2025). Confidence-Calibrated RAG. arXiv preprint.",
    "Pellissier, R. et Nenzhelele, T. E. (2013). Towards a Universal Competitive Intelligence Process Model. South African Journal of Information Management, 15(2), 1-7.",
    "Porter, M. E. (1985). Competitive Advantage: Creating and Sustaining Superior Performance. New York : Free Press.",
    "Porter, M. E. (2008). The Five Competitive Forces That Shape Strategy. Harvard Business Review, 86(1), 78-93.",
    "Rouach, D. et Santi, P. (2001). Competitive Intelligence Adds Value: Five Intelligence Attitudes. European Management Journal, 19(5), 552-559.",
    "Sahut, J.-M. et al. (2021). Applications of Text Mining in Services Management: A Systematic Literature Review. International Journal of Information Management Data Insights, 1(1), 100001.",
    "Singh, A. et Ehtesham, A. (2025). Agentic Retrieval-Augmented Generation: A Survey on Agentic RAG. arXiv preprint.",
    "Springer (2025). Retrieval-Augmented Generation (RAG). Business & Information Systems Engineering.",
    "Sridhara, S. et al. (2024). FACTS About Building Retrieval Augmented Generation-based Chatbots. arXiv preprint arXiv:2407.07858.",
    "Trim, P. R. J. et Lee, Y.-I. (2008). A Strategic Approach to Sustainable Partnership Development. European Business Review, 20(3), 222-237.",
    "Vaswani, A. et al. (2017). Attention Is All You Need. Advances in Neural Information Processing Systems, 30, 5998-6008.",
    "Veturi et al. (2024). Enhancing LLM Performance with Organizational Data via RAG. In Springer Business & Information Systems Engineering.",
    "Ye, Z. et al. (2024). R2AG: Incorporating Retrieval Information into RAG. arXiv preprint.",
]

for ref in references:
    p = doc.add_paragraph(ref)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)

# ── Save ──
output_path = "/home/yoyo/clawd/projects/memoire/revue-litterature.docx"
doc.save(output_path)
print(f"Document saved: {output_path}")
